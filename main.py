import openai
from docx import Document
from docx.shared import RGBColor
import pandas as pd
import json
import re
import time

# إعداد OpenAI
client = openai.OpenAI(api_key="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")

# الملفات
INPUT_FILE = "‏‏my_file.docx"
OUTPUT_FILE = "output.docx"
EXCEL_FILE = "entities_output.xlsx"

doc = Document(INPUT_FILE)

# أنماط الكيانات
style_map = {
    "الأماكن": {"style": "أماكن", "color": RGBColor(255, 0, 0)},
    "الأعلام": {"style": "أعلام", "color": RGBColor(0, 0, 255)},
    "الفرق": {"style": "فرق", "color": RGBColor(0, 128, 0)},
    "الكتب": {"style": "كتب", "color": RGBColor(128, 0, 128)},
    "الشواهد": {"style": "شواهد", "color": RGBColor(255, 165, 0)},
}

# توليد عنوان
def generate_title(paragraph_text):
    prompt = f"""
اقرأ النص التالي واستخرج منه عنوانًا موجزًا يدل على فكرته الأساسية. لا تشرح، فقط أعطني العنوان نفسه دون علامات تنصيص.
الرجاء الرد فقط بالعنوان دون أي مقدمة أو تعليق.
النص:
{paragraph_text}
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4",  # استخدم gpt-4 لو رغبت
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
        )
        return response.choices[0].message.content.strip().strip('"“”')
    except Exception as e:
        print("❌ خطأ في توليد العنوان:", e)
        return "عنوان غير متوفر"

# استخراج الكيانات
def extract_entities(paragraph_text):
    prompt = f"""
           أنت خبير لغوي مكلف بتحليل النص التالي واستخراج الكيانات التالية بدقة:

           1. الأماكن (المدن، الدول، المناطق)
              - ❌ لا تدرج أسماء منسوبة (مثل: البغداديين)
              - ✅ أدرج فقط الاسم الأصلي مثل: "بغداد"، "مكة"
              بمعنى أنك لو وجدت كلمة تعني نسبة شخص لمنطقة فلا تعتبرها مكاناً مثلا البغدادي لا يعني بغداد
              فقط ضع اعتبار للمكان المجرد وأعتبره من الأماكن سواء كان مباشراً مثل يمن أو معرفاً مثل اليمن لكن لو كان نسبة مثل اليمني فلا تضفه

           2. الأعلام (أسماء أشخاص حقيقيين فقط)
              - ❌ لا تدرج ألقاب جماعية أو صفات نسبية مثل: "الشافعية"، "الحنابلة"
              - ✅ فقط الأسماء الصريحة مثل: "ابن تيمية"، "أبو حامد الغزالي"

           3. الفرق والمذاهب (الكيانات الدينية والفكرية المعروفة)
              - ✅ مثل: "المرجئة"، "الأشاعرة"، "المعتزلة"

           4. أسماء الكتب (عناوين كتب علمية أو شرعية)
              - ✅ مثل: "صحيح البخاري"، "الإحياء"، "الرد على المنطقيين"

           5. الشواهد (آيات أو أحاديث فقط)
              - ✅ كل ما يبدأ بـ "قال الله تعالى" أو "قال رسول الله"

           ❗ المطلوب: استخراج الكيانات بدقة وتجاهل أي كلمات غير مطابقة للمعايير السابقة.

           النص:
{paragraph_text}

الرجاء الرد بصيغة JSON فقط كما يلي:
{{
  "الأماكن": [],
  "الأعلام": [],
  "الفرق": [],
  "الكتب": [],
  "الشواهد": []
}}
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
        )
        raw_content = response.choices[0].message.content
        cleaned = raw_content.replace("،", ",").replace("“", "\"").replace("”", "\"")
        match = re.search(r'\{[\s\S]*\}', cleaned)
        if match:
            json_part = match.group()
            entities = json.loads(json_part)
            return clean_entities(entities)
        else:
            print("❌ لم يتم العثور على JSON في الرد")
            return None
    except Exception as e:
        print("❌ خطأ أثناء تحليل GPT:", e)
        return None

# تنظيف الكيانات
def clean_entities(raw_entities):
    cleaned = {}
    for entity_type, items in raw_entities.items():
        final = []
        for item in items:
            if isinstance(item, str) and len(item.strip()) > 1:
                final.append(item.strip())
        cleaned[entity_type] = final
    return cleaned

# تنسيق الكيانات داخل الفقرة
def rewrite_paragraph_with_styles(paragraph, entities):
    full_text = paragraph.text
    if "\n" in full_text:
        title_line, body_text = full_text.split("\n", 1)
    else:
        title_line = ""
        body_text = full_text

    new_runs = []
    pointer = 0
    matches = []

    for entity_type, values in entities.items():
        for val in values:
            start = body_text.find(val, pointer)
            if start != -1:
                matches.append((start, start + len(val), val, entity_type))

    matches.sort()
    for start, end, val, etype in matches:
        if start > pointer:
            new_runs.append(("plain", body_text[pointer:start]))
        new_runs.append((etype, val))
        pointer = end
    if pointer < len(body_text):
        new_runs.append(("plain", body_text[pointer:]))

    # حذف المحتوى الأصلي
    for run in paragraph.runs:
        run.text = ""

    # إعادة إدخال العنوان في البداية
    title_run = paragraph.add_run(f"{title_line}\n\n")
    title_run.bold = True

    # إدخال المحتوى المنسق
    for style_type, content in new_runs:
        run = paragraph.add_run(content)
        if style_type != "plain":
            run.style = style_map[style_type]["style"]
            run.font.color.rgb = style_map[style_type]["color"]

# قائمة الكيانات
all_entities = []

# تنفيذ المعالجة
for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue

    print(f"🔎 تحليل فقرة: {text[:40]}...")

    # توليد العنوان
    title = generate_title(text)

    # إنشاء نص جديد يضم العنوان ثم النص الأصلي
    full_text = f"{title}\n{text}"

    # استبدال محتوى الفقرة بالنص الكامل
    for run in para.runs:
        run.text = ""
    para.add_run(full_text)

    # استخراج الكيانات من النص فقط (بدون العنوان)
    entities = extract_entities(text)
    all_entities.append({"الكيان": f"[عنوان] {title}", "النوع": "عنوان", "الفقرة": text})

    if entities:
        rewrite_paragraph_with_styles(para, entities)
        for etype, vals in entities.items():
            for v in vals:
                all_entities.append({"الكيان": v, "النوع": etype, "الفقرة": text})

    time.sleep(1.5)

# حفظ ملف Word النهائي
doc.save(OUTPUT_FILE)
print(f"✅ تم حفظ الملف المعدل: {OUTPUT_FILE}")

# حفظ الكيانات في Excel
df = pd.DataFrame(all_entities)
df.to_excel(EXCEL_FILE, index=False)
print(f"📊 تم حفظ الكيانات في: {EXCEL_FILE}")
