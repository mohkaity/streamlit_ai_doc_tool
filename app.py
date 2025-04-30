import streamlit as st
import openai
import pandas as pd
from docx import Document
from docx.shared import RGBColor
import tempfile
import json
import re
import time
import os

# --- Ø¥Ø¹Ø¯Ø§Ø¯ Ø§Ù„ÙˆØ§Ø¬Ù‡Ø© ---
st.set_page_config(page_title="ØªØ­Ù„ÙŠÙ„ Ø§Ù„Ù…Ø³ØªÙ†Ø¯ Ø¨Ø§Ù„Ø°ÙƒØ§Ø¡ Ø§Ù„Ø§ØµØ·Ù†Ø§Ø¹ÙŠ", layout="centered")
st.title("ðŸ“„ Ø£Ø¯Ø§Ø© ØªØ­Ù„ÙŠÙ„ Ø§Ù„ÙÙ‚Ø±Ø§Øª ÙˆØ§Ø³ØªØ®Ø±Ø§Ø¬ Ø§Ù„ÙƒÙŠØ§Ù†Ø§Øª")

# --- Ø¥Ø¯Ø®Ø§Ù„ Ø§Ù„Ø¨ÙŠØ§Ù†Ø§Øª ---
st.sidebar.header("Ø¥Ø¹Ø¯Ø§Ø¯Ø§Øª Ø§Ù„Ù†Ù…ÙˆØ°Ø¬")
api_key = st.sidebar.text_input("ðŸ”‘ Ø£Ø¯Ø®Ù„ Ù…ÙØªØ§Ø­ OpenAI API", type="password")
model = st.sidebar.selectbox("ðŸ¤– Ø§Ø®ØªØ± Ø§Ù„Ù†Ù…ÙˆØ°Ø¬", ["gpt-3.5-turbo", "gpt-4"])

uploaded_file = st.file_uploader("ðŸ“¤ Ø§Ø±ÙØ¹ Ù…Ù„Ù Word (Ø¨ØµÙŠØºØ© .docx)", type=["docx"])

# --- Ø¥Ø¹Ø¯Ø§Ø¯ Ø§Ù„Ø£Ù†Ù…Ø§Ø· ---
style_map = {
    "Ø§Ù„Ø£Ù…Ø§ÙƒÙ†": {"style": "Ø£Ù…Ø§ÙƒÙ†", "color": RGBColor(255, 0, 0)},
    "Ø§Ù„Ø£Ø¹Ù„Ø§Ù…": {"style": "Ø£Ø¹Ù„Ø§Ù…", "color": RGBColor(0, 0, 255)},
    "Ø§Ù„ÙØ±Ù‚": {"style": "ÙØ±Ù‚", "color": RGBColor(0, 128, 0)},
    "Ø§Ù„ÙƒØªØ¨": {"style": "ÙƒØªØ¨", "color": RGBColor(128, 0, 128)},
    "Ø§Ù„Ø´ÙˆØ§Ù‡Ø¯": {"style": "Ø´ÙˆØ§Ù‡Ø¯", "color": RGBColor(255, 165, 0)},
}

# --- Ø¯ÙˆØ§Ù„ GPT ---
def generate_title(client, model, paragraph_text):
    prompt = f"""
Ø§Ù‚Ø±Ø£ Ø§Ù„Ù†Øµ Ø§Ù„ØªØ§Ù„ÙŠ ÙˆØ§Ø³ØªØ®Ø±Ø¬ Ù…Ù†Ù‡ Ø¹Ù†ÙˆØ§Ù†Ù‹Ø§ Ù…ÙˆØ¬Ø²Ù‹Ø§ ÙŠØ¯Ù„ Ø¹Ù„Ù‰ ÙÙƒØ±ØªÙ‡ Ø§Ù„Ø£Ø³Ø§Ø³ÙŠØ©. Ù„Ø§ ØªØ´Ø±Ø­ØŒ ÙÙ‚Ø· Ø£Ø¹Ø·Ù†ÙŠ Ø§Ù„Ø¹Ù†ÙˆØ§Ù† Ù†ÙØ³Ù‡ Ø¯ÙˆÙ† Ø¹Ù„Ø§Ù…Ø§Øª ØªÙ†ØµÙŠØµ.
Ø§Ù„Ø±Ø¬Ø§Ø¡ Ø§Ù„Ø±Ø¯ ÙÙ‚Ø· Ø¨Ø§Ù„Ø¹Ù†ÙˆØ§Ù† Ø¯ÙˆÙ† Ø£ÙŠ Ù…Ù‚Ø¯Ù…Ø© Ø£Ùˆ ØªØ¹Ù„ÙŠÙ‚.
Ø§Ù„Ù†Øµ:
{paragraph_text}
"""
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )
    return response.choices[0].message.content.strip().strip('"â€œâ€')

def extract_entities(client, model, paragraph_text):
    prompt = f"""
    Ø£Ù†Øª Ø®Ø¨ÙŠØ± Ù„ØºÙˆÙŠ Ù…ÙƒÙ„Ù Ø¨ØªØ­Ù„ÙŠÙ„ Ø§Ù„Ù†Øµ Ø§Ù„ØªØ§Ù„ÙŠ ÙˆØ§Ø³ØªØ®Ø±Ø§Ø¬ Ø§Ù„ÙƒÙŠØ§Ù†Ø§Øª Ø§Ù„ØªØ§Ù„ÙŠØ© Ø¨Ø¯Ù‚Ø©:

           1. Ø§Ù„Ø£Ù…Ø§ÙƒÙ† (Ø§Ù„Ù…Ø¯Ù†ØŒ Ø§Ù„Ø¯ÙˆÙ„ØŒ Ø§Ù„Ù…Ù†Ø§Ø·Ù‚)
              - âŒ Ù„Ø§ ØªØ¯Ø±Ø¬ Ø£Ø³Ù…Ø§Ø¡ Ù…Ù†Ø³ÙˆØ¨Ø© (Ù…Ø«Ù„: Ø§Ù„Ø¨ØºØ¯Ø§Ø¯ÙŠÙŠÙ†)
              - âœ… Ø£Ø¯Ø±Ø¬ ÙÙ‚Ø· Ø§Ù„Ø§Ø³Ù… Ø§Ù„Ø£ØµÙ„ÙŠ Ù…Ø«Ù„: "Ø¨ØºØ¯Ø§Ø¯"ØŒ "Ù…ÙƒØ©"
              Ø¨Ù…Ø¹Ù†Ù‰ Ø£Ù†Ùƒ Ù„Ùˆ ÙˆØ¬Ø¯Øª ÙƒÙ„Ù…Ø© ØªØ¹Ù†ÙŠ Ù†Ø³Ø¨Ø© Ø´Ø®Øµ Ù„Ù…Ù†Ø·Ù‚Ø© ÙÙ„Ø§ ØªØ¹ØªØ¨Ø±Ù‡Ø§ Ù…ÙƒØ§Ù†Ø§Ù‹ Ù…Ø«Ù„Ø§ Ø§Ù„Ø¨ØºØ¯Ø§Ø¯ÙŠ Ù„Ø§ ÙŠØ¹Ù†ÙŠ Ø¨ØºØ¯Ø§Ø¯
              ÙÙ‚Ø· Ø¶Ø¹ Ø§Ø¹ØªØ¨Ø§Ø± Ù„Ù„Ù…ÙƒØ§Ù† Ø§Ù„Ù…Ø¬Ø±Ø¯ ÙˆØ£Ø¹ØªØ¨Ø±Ù‡ Ù…Ù† Ø§Ù„Ø£Ù…Ø§ÙƒÙ† Ø³ÙˆØ§Ø¡ ÙƒØ§Ù† Ù…Ø¨Ø§Ø´Ø±Ø§Ù‹ Ù…Ø«Ù„ ÙŠÙ…Ù† Ø£Ùˆ Ù…Ø¹Ø±ÙØ§Ù‹ Ù…Ø«Ù„ Ø§Ù„ÙŠÙ…Ù† Ù„ÙƒÙ† Ù„Ùˆ ÙƒØ§Ù† Ù†Ø³Ø¨Ø© Ù…Ø«Ù„ Ø§Ù„ÙŠÙ…Ù†ÙŠ ÙÙ„Ø§ ØªØ¶ÙÙ‡

           2. Ø§Ù„Ø£Ø¹Ù„Ø§Ù… (Ø£Ø³Ù…Ø§Ø¡ Ø£Ø´Ø®Ø§Øµ Ø­Ù‚ÙŠÙ‚ÙŠÙŠÙ† ÙÙ‚Ø·)
              - âŒ Ù„Ø§ ØªØ¯Ø±Ø¬ Ø£Ù„Ù‚Ø§Ø¨ Ø¬Ù…Ø§Ø¹ÙŠØ© Ø£Ùˆ ØµÙØ§Øª Ù†Ø³Ø¨ÙŠØ© Ù…Ø«Ù„: "Ø§Ù„Ø´Ø§ÙØ¹ÙŠØ©"ØŒ "Ø§Ù„Ø­Ù†Ø§Ø¨Ù„Ø©"
              - âœ… ÙÙ‚Ø· Ø§Ù„Ø£Ø³Ù…Ø§Ø¡ Ø§Ù„ØµØ±ÙŠØ­Ø© Ù…Ø«Ù„: "Ø§Ø¨Ù† ØªÙŠÙ…ÙŠØ©"ØŒ "Ø£Ø¨Ùˆ Ø­Ø§Ù…Ø¯ Ø§Ù„ØºØ²Ø§Ù„ÙŠ"

           3. Ø§Ù„ÙØ±Ù‚ ÙˆØ§Ù„Ù…Ø°Ø§Ù‡Ø¨ (Ø§Ù„ÙƒÙŠØ§Ù†Ø§Øª Ø§Ù„Ø¯ÙŠÙ†ÙŠØ© ÙˆØ§Ù„ÙÙƒØ±ÙŠØ© Ø§Ù„Ù…Ø¹Ø±ÙˆÙØ©)
              - âœ… Ù…Ø«Ù„: "Ø§Ù„Ù…Ø±Ø¬Ø¦Ø©"ØŒ "Ø§Ù„Ø£Ø´Ø§Ø¹Ø±Ø©"ØŒ "Ø§Ù„Ù…Ø¹ØªØ²Ù„Ø©"

           4. Ø£Ø³Ù…Ø§Ø¡ Ø§Ù„ÙƒØªØ¨ (Ø¹Ù†Ø§ÙˆÙŠÙ† ÙƒØªØ¨ Ø¹Ù„Ù…ÙŠØ© Ø£Ùˆ Ø´Ø±Ø¹ÙŠØ©)
              - âœ… Ù…Ø«Ù„: "ØµØ­ÙŠØ­ Ø§Ù„Ø¨Ø®Ø§Ø±ÙŠ"ØŒ "Ø§Ù„Ø¥Ø­ÙŠØ§Ø¡"ØŒ "Ø§Ù„Ø±Ø¯ Ø¹Ù„Ù‰ Ø§Ù„Ù…Ù†Ø·Ù‚ÙŠÙŠÙ†"

           5. Ø§Ù„Ø´ÙˆØ§Ù‡Ø¯ (Ø¢ÙŠØ§Øª Ø£Ùˆ Ø£Ø­Ø§Ø¯ÙŠØ« ÙÙ‚Ø·)
              - âœ… ÙƒÙ„ Ù…Ø§ ÙŠØ¨Ø¯Ø£ Ø¨Ù€ "Ù‚Ø§Ù„ Ø§Ù„Ù„Ù‡ ØªØ¹Ø§Ù„Ù‰" Ø£Ùˆ "Ù‚Ø§Ù„ Ø±Ø³ÙˆÙ„ Ø§Ù„Ù„Ù‡"

           â— Ø§Ù„Ù…Ø·Ù„ÙˆØ¨: Ø§Ø³ØªØ®Ø±Ø§Ø¬ Ø§Ù„ÙƒÙŠØ§Ù†Ø§Øª Ø¨Ø¯Ù‚Ø© ÙˆØªØ¬Ø§Ù‡Ù„ Ø£ÙŠ ÙƒÙ„Ù…Ø§Øª ØºÙŠØ± Ù…Ø·Ø§Ø¨Ù‚Ø© Ù„Ù„Ù…Ø¹Ø§ÙŠÙŠØ± Ø§Ù„Ø³Ø§Ø¨Ù‚Ø©.

           Ø§Ù„Ù†Øµ:
{paragraph_text}

Ø§Ù„Ø±Ø¬Ø§Ø¡ Ø§Ù„Ø±Ø¯ Ø¨ØµÙŠØºØ© JSON ÙÙ‚Ø· ÙƒÙ…Ø§ ÙŠÙ„ÙŠ:
{{
  "Ø§Ù„Ø£Ù…Ø§ÙƒÙ†": [],
  "Ø§Ù„Ø£Ø¹Ù„Ø§Ù…": [],
  "Ø§Ù„ÙØ±Ù‚": [],
  "Ø§Ù„ÙƒØªØ¨": [],
  "Ø§Ù„Ø´ÙˆØ§Ù‡Ø¯": []
}}
"""
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
    )
    raw_content = response.choices[0].message.content
    cleaned = raw_content.replace("ØŒ", ",").replace("â€œ", '"').replace("â€", '"')
    match = re.search(r'\{[\s\S]*\}', cleaned)
    if match:
        return json.loads(match.group())
    else:
        return {}

# --- ØªÙ†Ø³ÙŠÙ‚ Ø§Ù„ÙƒÙŠØ§Ù†Ø§Øª ---
def rewrite_paragraph_with_styles(paragraph, entities):
    full_text = paragraph.text
    new_runs = []
    pointer = 0
    matches = []

    for entity_type, values in entities.items():
        for val in values:
            start = full_text.find(val, pointer)
            if start != -1:
                matches.append((start, start + len(val), val, entity_type))

    matches.sort()
    for start, end, val, etype in matches:
        if start > pointer:
            new_runs.append(("plain", full_text[pointer:start]))
        new_runs.append((etype, val))
        pointer = end
    if pointer < len(full_text):
        new_runs.append(("plain", full_text[pointer:]))

    for run in paragraph.runs:
        run.text = ""

    for style_type, content in new_runs:
        run = paragraph.add_run(content)
        if style_type != "plain":
            run.style = style_map[style_type]["style"]
            run.font.color.rgb = style_map[style_type]["color"]

# --- Ø§Ù„Ù…Ø¹Ø§Ù„Ø¬Ø© ---
if uploaded_file and api_key:
    with st.spinner("ðŸ” Ø¬Ø§Ø±ÙŠ ØªØ­Ù„ÙŠÙ„ Ø§Ù„Ù…Ù„Ù..."):
        client = openai.OpenAI(api_key=api_key)
        doc = Document(uploaded_file)
        new_doc = Document()
        all_entities = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            title = generate_title(client, model, text)
            title_para = new_doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True

            body_para = new_doc.add_paragraph(text)
            entities = extract_entities(client, model, text)
            all_entities.append({"Ø§Ù„ÙƒÙŠØ§Ù†": f"[Ø¹Ù†ÙˆØ§Ù†] {title}", "Ø§Ù„Ù†ÙˆØ¹": "Ø¹Ù†ÙˆØ§Ù†", "Ø§Ù„ÙÙ‚Ø±Ø©": text})

            if entities:
                rewrite_paragraph_with_styles(body_para, entities)
                for etype, vals in entities.items():
                    for v in vals:
                        all_entities.append({"Ø§Ù„ÙƒÙŠØ§Ù†": v, "Ø§Ù„Ù†ÙˆØ¹": etype, "Ø§Ù„ÙÙ‚Ø±Ø©": text})

            time.sleep(1.5)

        # --- Ø­ÙØ¸ Ø§Ù„Ù…Ù„ÙØ§Øª Ø§Ù„Ù…Ø¤Ù‚ØªØ© ---
        word_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        new_doc.save(word_file.name)

        df = pd.DataFrame(all_entities)
        excel_file = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
        df.to_excel(excel_file.name, index=False)

        st.success("âœ… ØªÙ… Ø§Ù„ØªØ­Ù„ÙŠÙ„ Ø¨Ù†Ø¬Ø§Ø­")
        st.download_button("ðŸ“¥ ØªØ­Ù…ÙŠÙ„ Ù…Ù„Ù Word Ø§Ù„Ù†Ø§ØªØ¬", data=open(word_file.name, "rb"), file_name="Ù†ØªÙŠØ¬Ø©.docx")
        st.download_button("ðŸ“Š ØªØ­Ù…ÙŠÙ„ Ù…Ù„Ù Excel Ù„Ù„ÙƒÙŠØ§Ù†Ø§Øª", data=open(excel_file.name, "rb"), file_name="Ø§Ù„ÙƒÙŠØ§Ù†Ø§Øª.xlsx")

elif uploaded_file and not api_key:
    st.warning("âš ï¸ Ø§Ù„Ø±Ø¬Ø§Ø¡ Ø¥Ø¯Ø®Ø§Ù„ Ù…ÙØªØ§Ø­ API Ù„Ù„Ù…ØªØ§Ø¨Ø¹Ø©.")
